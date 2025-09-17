# -*- coding: utf-8 -*-
"""
Excel → mikro-tabele w Wordzie (GUI, Tkinter) z mapowaniem, transformacjami i presetami JSON.
UX:
- Po wybraniu pliku Excel kolumny wczytują się automatycznie (brak osobnego przycisku).
- „Pozostałe kolumny” pozostają puste do momentu wczytania presetu (zamiast listować wszystko).
- Opcje: ramki foto/mapa, przecinek dziesiętny, podział strony, marginesy, sortowanie.
"""

# ---------- BOOTSTRAP PIP ----------
import sys, subprocess
def _pip_install(spec):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-U", spec])
    except Exception:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "-U", spec])

def _ensure_min(pkg, module, min_v):
    try:
        try:
            from importlib.metadata import version as _v
        except Exception:
            from importlib_metadata import version as _v
        cur = _v(module)
    except Exception:
        cur = None
    def _lt(a,b):
        pa=lambda s:[int(x) for x in str(s).split(".") if x.isdigit()]
        return pa(a or "0")<pa(b)
    if cur is None or _lt(cur, min_v):
        _pip_install(f"{pkg}>={min_v}")

def _ensure(pkg, module=None):
    import importlib.util
    if not importlib.util.find_spec(module or pkg.replace("-","_")):
        _pip_install(pkg)

_ensure("pandas","pandas"); _ensure("python-docx","docx"); _ensure("openpyxl","openpyxl"); _ensure_min("xlrd","xlrd","2.0.1")

# ---------- IMPORTY ----------
import os, json
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date, datetime

EMU_PER_CM = 360000  # Word EMU -> cm

# ---------- TRANSFORMACJE ----------
def _is_missing(v):
    return v is None or (isinstance(v, float) and pd.isna(v)) or (isinstance(v, str) and v.strip()=="")
def tf_identity(val, row, comma=False):
    return "" if _is_missing(val) else str(val)
def tf_m2_to_ha_round2(val, row, comma=False):
    if _is_missing(val): return ""
    try:
        ha = round(float(str(val).replace(",", "."))/10000.0, 2)
        s = f"{ha:.2f}"
        return s.replace(".", ",") if comma else s
    except Exception:
        return ""
def tf_prelim_to_bedomning(val, row, comma=False):
    v = "" if _is_missing(val) else str(val).strip().lower()
    return "Säker" if v in {"0","nej","no","false","0.0",""} else "Preliminärt"
def tf_constant(val, row, comma=False, const_val=""):
    return const_val
def tf_date_only(val, row, comma=False):
    if _is_missing(val): return ""
    try:
        if isinstance(val, (datetime, date, pd.Timestamp)):
            return pd.to_datetime(val).date().isoformat()
        s = str(val).strip()
        if " " in s: s = s.split(" ")[0]
        if "T" in s: s = s.split("T")[0]
        d = pd.to_datetime(s, errors="coerce")
        return d.date().isoformat() if not pd.isna(d) else s
    except Exception:
        return str(val)
def tf_bool_ja_nej(val, row, comma=False):
    """1/0, 1.0/0.0, ja/nej, yes/no, true/false → 'Ja'/'Nej'; nieznane → '' (później ' - ')."""
    if _is_missing(val): return ""
    s = str(val).strip().lower()
    try:
        f = float(s.replace(",", "."))
        return "Ja" if f != 0.0 else "Nej"
    except Exception:
        pass
    positives = {"ja","yes","y","true","t","1"}
    negatives = {"nej","no","n","false","f","0"}
    if s in positives: return "Ja"
    if s in negatives: return "Nej"
    return ""

TRANSFORMS = {
    "identity": ("Bez zmian", tf_identity),
    "m2_to_ha_round2": ("m² → ha (0,01)", tf_m2_to_ha_round2),
    "prelim_to_bedomning": ("0/nej → Säker, inne → Preliminärt", tf_prelim_to_bedomning),
    "constant": ("Stała wartość", tf_constant),
    "date_only": ("Tylko data (YYYY-MM-DD)", tf_date_only),
    "bool_ja_nej": ("Bool → Ja/Nej", tf_bool_ja_nej),
}

# ---------- DOCX HELPERS ----------
def _shade_cell(cell, fill_hex="F2F2F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)
def _set_cell_text(cell, text, bold=False, size_pt=10):
    cell.text = "" if text is None else str(text)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = bold; r.font.size = Pt(size_pt)
    if cell.paragraphs: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------- EXCEL ----------
def read_excel_any(path):
    ext = os.path.splitext(path.lower())[1]
    if ext == ".xls": return pd.read_excel(path, engine="xlrd")
    return pd.read_excel(path)

# ---------- PRESETY: dopasowanie aliasów ----------
def _norm(s: str) -> str:
    return "".join(ch for ch in str(s).lower() if ch.isalnum())
def resolve_source_name(wanted: str, df_columns, aliases: dict) -> str:
    cols = list(map(str, df_columns))
    by_norm = {_norm(c): c for c in cols}
    candidates = [wanted] + list(aliases.get(wanted, [])) if wanted else []
    for cand in candidates:
        n = _norm(cand)
        if n in by_norm: return by_norm[n]
        for c in cols:
            if n and n in _norm(c): return c
    return ""  # brak trafienia

# ---------- GENERACJA DOCX ----------
def build_docx(df, mapping_rows, extra_cols, out_path, out_name,
               add_photo=True, photo_h_cm=6.0,
               add_map=True, map_h_cm=6.0,
               page_break=False, use_decimal_comma=True,
               margin_left_cm=2.0, margin_right_cm=2.0,
               margin_top_cm=2.0, margin_bottom_cm=2.0,
               layout_mode="A"):
    """
    layout_mode: 'A' albo 'B'
      A: tabela, potem (foto?), potem (mapa?)
      B: tabela po lewej + scalona kolumna na zdjęcie po prawej; mapa (opcjonalnie) pod spodem.
         Jeśli add_photo=False -> zachowuje się jak A (bez kolumny zdjęcia).
    """
    doc = Document()
    section = doc.sections[0]
    # Marginesy
    section.left_margin   = Cm(float(margin_left_cm))
    section.right_margin  = Cm(float(margin_right_cm))
    section.top_margin    = Cm(float(margin_top_cm))
    section.bottom_margin = Cm(float(margin_bottom_cm))

    # Szerokość treści po marginesach (w cm)
    content_w_cm = (section.page_width - section.left_margin - section.right_margin) / EMU_PER_CM

    # --- Funkcja pomocnicza: zbuduj klasyczną tabelę 2 kol. (etykieta/wartość) ---
    def build_info_table(parent_container, rows, col0_cm, col1_cm):
        t = parent_container.add_table(rows=0, cols=2) if hasattr(parent_container, "add_table") else doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = "Table Grid"; t.autofit = False
        for item in rows:
            if not item.get("enabled", False): continue
            lbl = item.get("label",""); src = item.get("source","")
            tf_key = item.get("transform","identity"); const_val = item.get("const","")
            raw = cur_row.get(src, "") if src else ""
            tf = TRANSFORMS.get(tf_key, TRANSFORMS["identity"])[1]
            val = tf(raw, cur_row, use_decimal_comma, const_val=const_val) if tf_key=="constant" else tf(raw, cur_row, use_decimal_comma)
            if tf_key != "constant" and (val is None or (isinstance(val,str) and val.strip()=="")):
                val = " - "
            cells = t.add_row().cells
            _set_cell_text(cells[0], lbl, bold=True); _shade_cell(cells[0], "EAEAEA")
            _set_cell_text(cells[1], val)
        for r in t.rows:
            r.cells[0].width = Cm(col0_cm); r.cells[1].width = Cm(col1_cm)
        return t

    # --- Dołącz pozostałe kolumny jako identity na końcu mapowania ---
    for col_name in extra_cols:
        mapping_rows.append({"enabled": True, "label": col_name, "source": col_name, "transform": "identity", "const": ""})

    # --- Generacja dla każdego rekordu ---
    for _, cur_row in df.iterrows():
        if layout_mode == "B" and add_photo:
            # Domyślny podział szerokości: prawa kolumna na foto ok. 35% szerokości
            right_photo_cm = max(5.0, min(8.5, content_w_cm * 0.35))
            left_outer_cm  = max(6.0, content_w_cm - right_photo_cm)

            # Tabela zewnętrzna 2 kolumny: [info][foto]
            outer = doc.add_table(rows=1, cols=2)
            outer.style = "Table Grid"; outer.autofit = False
            outer.alignment = WD_TABLE_ALIGNMENT.LEFT
            # szerokości kolumn zewnętrznych
            outer.rows[0].cells[0].width = Cm(left_outer_cm)
            outer.rows[0].cells[1].width = Cm(right_photo_cm)

            left_cell  = outer.cell(0,0)
            right_cell = outer.cell(0,1)

            # W LEWYM: tabela informacji (2 kolumny). Udział etykiety ~45% lewego bloku
            label_cm = max(3.5, min(7.0, left_outer_cm * 0.45))
            value_cm = max(3.5, left_outer_cm - label_cm)
            build_info_table(left_cell, mapping_rows, label_cm, value_cm)

            # W PRAWYM: podpis + ramka foto
            if add_photo:
                p = right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph("")
                run = p.add_run("Representativt foto:")
                run.font.size = Pt(10)
                # Jednokomórkowa ramka wewnątrz prawej komórki
                ph = right_cell.add_table(rows=1, cols=1)
                ph.style = "Table Grid"; ph.autofit = False
                ph.rows[0].cells[0].width = Cm(right_photo_cm)
                # Wysokość ramki: przez wysokość wiersza
                ph.rows[0].height = Cm(float(photo_h_cm))
                _set_cell_text(ph.cell(0, 0), " ")
            doc.add_paragraph("")
            # Mapa na dole (jeśli włączona)
            if add_map:
                p2 = doc.add_paragraph("Kartbild av objektet:")
                if p2.runs: p2.runs[0].font.size = Pt(10)
                mp = doc.add_table(rows=1, cols=1); mp.style = "Table Grid"; mp.autofit = False
                mp.rows[0].height = Cm(map_h_cm); mp.cell(0, 0).width = Cm(content_w_cm)
                _set_cell_text(mp.cell(0, 0), " ")
                doc.add_paragraph("")

        else:
            # Układ A (lub fallback, jeśli w B brak zdjęcia)
            # Klasyczna tabela 2 kolumny
            left_w_cm  = 6.0
            right_w_cm = max(6.0, content_w_cm - left_w_cm)
            build_info_table(doc, mapping_rows, left_w_cm, right_w_cm)
            doc.add_paragraph("")
            if add_photo:
                p = doc.add_paragraph("Representativt foto:")
                if p.runs: p.runs[0].font.size = Pt(10)
                ph = doc.add_table(rows=1, cols=1); ph.style = "Table Grid"; ph.autofit = False
                ph.rows[0].height = Cm(photo_h_cm); ph.cell(0, 0).width = Cm(content_w_cm)
                _set_cell_text(ph.cell(0, 0), " ")
                doc.add_paragraph("")
            if add_map:
                p2 = doc.add_paragraph("Kartbild av objektet:")
                if p2.runs: p2.runs[0].font.size = Pt(10)
                mp = doc.add_table(rows=1, cols=1); mp.style = "Table Grid"; mp.autofit = False
                mp.rows[0].height = Cm(map_h_cm); mp.cell(0, 0).width = Cm(content_w_cm)
                _set_cell_text(mp.cell(0, 0), " ")
                doc.add_paragraph("")

        if page_break:
            doc.add_page_break()

    out_file = os.path.join(out_path, out_name if out_name.lower().endswith(".docx") else out_name + ".docx")
    doc.save(out_file); return out_file

# ---------- GUI ----------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Excel → Mikro-tabele Word (presety) — układ A/B")
        self.geometry("1050x900"); self.resizable(True, True)
        self.df = None
        self.mapping = []
        self.extra_vars = {}   # nazwa kolumny -> tk.BooleanVar()
        self.preset_label_var = tk.StringVar(value="(brak)")
        self._build_ui()

    def _build_ui(self):
        pad = {"padx":8, "pady":6}
        top = ttk.Frame(self); top.pack(fill="both", expand=True)
        ttk.Label(top, text="Plik Excel (.xlsx / .xls):").grid(row=0, column=0, sticky="w", **pad)
        self.in_entry = ttk.Entry(top, width=80); self.in_entry.grid(row=0, column=1, sticky="we", **pad)
        ttk.Button(top, text="Wybierz…", command=self.pick_excel).grid(row=0, column=2, **pad)
        ttk.Label(top, text="Folder wyjściowy:").grid(row=1, column=0, sticky="w", **pad)
        self.out_entry = ttk.Entry(top, width=80); self.out_entry.grid(row=1, column=1, sticky="we", **pad)
        ttk.Button(top, text="Wybierz…", command=self.pick_outdir).grid(row=1, column=2, **pad)
        ttk.Label(top, text="Nazwa pliku .docx:").grid(row=2, column=0, sticky="w", **pad)
        self.name_entry = ttk.Entry(top, width=40); self.name_entry.insert(0, "wynik.docx"); self.name_entry.grid(row=2, column=1, sticky="w", **pad)

        btn_row = ttk.Frame(top); btn_row.grid(row=3, column=0, columnspan=3, sticky="we", padx=8, pady=(2,8))
        ttk.Button(btn_row, text="Wczytaj preset…", command=self.load_preset_dialog).pack(side="left")
        ttk.Label(btn_row, text="Preset:").pack(side="left", padx=(16,4))
        ttk.Label(btn_row, textvariable=self.preset_label_var, foreground="#555").pack(side="left")

        nb = ttk.Notebook(top); nb.grid(row=4, column=0, columnspan=3, sticky="nsew", padx=8, pady=8)
        top.rowconfigure(4, weight=1); top.columnconfigure(1, weight=1)

        # --- Mapowanie
        self.tab_map = ttk.Frame(nb); nb.add(self.tab_map, text="Mapowanie pól")
        self._build_map_tab(self.tab_map)

        # --- Pozostałe kolumny
        self.tab_extra = ttk.Frame(nb); nb.add(self.tab_extra, text="Pozostałe kolumny")
        self.extra_box = ttk.LabelFrame(self.tab_extra, text="Kolumny (nieużyte w mapie)")
        self.extra_box.pack(fill="both", expand=True, padx=8, pady=8)

        # --- Opcje
        self.tab_opt = ttk.Frame(nb); nb.add(self.tab_opt, text="Opcje")
        self._build_opt_tab(self.tab_opt)

        ttk.Button(top, text="Generuj DOCX", command=self.run)\
            .grid(row=5, column=0, columnspan=3, sticky="we", padx=8, pady=8)

    def _build_map_tab(self, parent):
        cols = ("order","#on","label","source","transform","const")
        self.tree = ttk.Treeview(parent, columns=cols, show="headings", height=12)
        for c, txt, w in [
            ("order","Kolej.",60),
            ("#on","Włącz",60),
            ("label","Nowa nazwa",260),
            ("source","Źródło",240),
            ("transform","Transformacja",240),
            ("const","Stała",200),
        ]:
            self.tree.heading(c, text=txt); self.tree.column(c, width=w, anchor="w")
        self.tree.pack(fill="both", expand=True, padx=8, pady=8)

        ed = ttk.LabelFrame(parent, text="Edytuj / dodaj wiersz")
        ed.pack(fill="x", padx=8, pady=(0,8))
        self.var_on = tk.BooleanVar(value=True)
        self.var_label = tk.StringVar(); self.var_source = tk.StringVar()
        self.var_transform = tk.StringVar(value="identity"); self.var_const = tk.StringVar()

        ttk.Checkbutton(ed, text="Włączony", variable=self.var_on).grid(row=0, column=0, sticky="w", padx=6, pady=4)
        ttk.Label(ed, text="Nowa nazwa:").grid(row=0, column=1, sticky="e", padx=6)
        self.ent_label = ttk.Entry(ed, textvariable=self.var_label, width=34); self.ent_label.grid(row=0, column=2, sticky="w", padx=6)

        ttk.Label(ed, text="Źródło:").grid(row=1, column=1, sticky="e", padx=6)
        self.cmb_source = ttk.Combobox(ed, textvariable=self.var_source, width=32, values=[], state="readonly")
        self.cmb_source.grid(row=1, column=2, sticky="w", padx=6)

        ttk.Label(ed, text="Transformacja:").grid(row=0, column=3, sticky="e", padx=6)
        self.cmb_transform = ttk.Combobox(ed, textvariable=self.var_transform, width=28,
                                          values=[k for k in TRANSFORMS], state="readonly")
        self.cmb_transform.grid(row=0, column=4, sticky="w", padx=6)

        ttk.Label(ed, text="Stała:").grid(row=1, column=3, sticky="e", padx=6)
        self.ent_const = ttk.Entry(ed, textvariable=self.var_const, width=30); self.ent_const.grid(row=1, column=4, sticky="w", padx=6)

        ttk.Button(ed, text="Zastosuj (aktualizuj wybrany)", command=self.apply_edit)\
            .grid(row=0, column=5, rowspan=2, sticky="nsw", padx=6, pady=4)

        btns = ttk.Frame(parent); btns.pack(fill="x", padx=8, pady=(0,4))
        ttk.Button(btns, text="↑", width=3, command=lambda: self.move_selected(-1)).pack(side="left", padx=2)
        ttk.Button(btns, text="↓", width=3, command=lambda: self.move_selected(1)).pack(side="left", padx=2)
        ttk.Button(btns, text="Usuń", command=self.delete_selected).pack(side="left", padx=8)
        ttk.Button(btns, text="Dodaj z kolumny…", command=self.add_from_column).pack(side="left", padx=8)
        ttk.Button(btns, text="Dodaj wiersz", command=self.add_blank_row).pack(side="left", padx=8)

        help_txt = ("Wybierz wiersz, edytuj pola poniżej i kliknij „Zastosuj”. "
                    "„Dodaj z kolumny…” tworzy wiersz mapujący istniejącą kolumnę. "
                    "„Dodaj wiersz” tworzy pusty wpis z transformacją 'constant' – uzupełnij pole „Stała”. "
                    "Brak wartości ze źródła daje „ - ” (nie dotyczy 'constant').")
        ttk.Label(parent, text=help_txt, foreground="#555", wraplength=900, justify="left")\
            .pack(fill="x", padx=12, pady=(0,8))

        self.tree.bind("<<TreeviewSelect>>", self.on_select_row)

    def _build_opt_tab(self, parent):
        opt = ttk.LabelFrame(parent, text="Opcje dokumentu")
        opt.pack(fill="x", padx=8, pady=8)
        self.var_photo = tk.BooleanVar(value=True); self.var_photo_h = tk.DoubleVar(value=6.0)
        self.var_map = tk.BooleanVar(value=True);   self.var_map_h = tk.DoubleVar(value=6.0)
        self.var_break = tk.BooleanVar(value=True); self.var_comma = tk.BooleanVar(value=True)

        # Marginesy [cm]
        self.var_marg_l = tk.DoubleVar(value=2.0)
        self.var_marg_r = tk.DoubleVar(value=2.0)
        self.var_marg_t = tk.DoubleVar(value=2.0)
        self.var_marg_b = tk.DoubleVar(value=2.0)

        ttk.Checkbutton(opt, text="Dodaj ramkę na zdjęcie po każdym rekordzie", variable=self.var_photo)\
            .grid(row=0, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(opt, text="Wysokość [cm]:").grid(row=0, column=1, sticky="e")
        ttk.Entry(opt, textvariable=self.var_photo_h, width=6).grid(row=0, column=2, sticky="w", padx=6)

        ttk.Checkbutton(opt, text="Dodaj ramkę na mapę (po zdjęciu/tabeli)", variable=self.var_map)\
            .grid(row=1, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(opt, text="Wysokość [cm]:").grid(row=1, column=1, sticky="e")
        ttk.Entry(opt, textvariable=self.var_map_h, width=6).grid(row=1, column=2, sticky="w", padx=6)

        ttk.Checkbutton(opt, text="Podział strony po każdym rekordzie", variable=self.var_break)\
            .grid(row=2, column=0, sticky="w", padx=8, pady=4)
        ttk.Checkbutton(opt, text="Użyj przecinka dziesiętnego", variable=self.var_comma)\
            .grid(row=2, column=1, sticky="w", padx=8, pady=4)

        # Marginesy [cm]
        ttk.Label(opt, text="Marginesy [cm]:").grid(row=3, column=0, sticky="w", padx=8, pady=(10,4))
        ttk.Label(opt, text="Lewy").grid(row=3, column=1, sticky="e");   ttk.Entry(opt, textvariable=self.var_marg_l, width=6).grid(row=3, column=2, sticky="w", padx=6)
        ttk.Label(opt, text="Prawy").grid(row=3, column=3, sticky="e");  ttk.Entry(opt, textvariable=self.var_marg_r, width=6).grid(row=3, column=4, sticky="w", padx=6)
        ttk.Label(opt, text="Górny").grid(row=4, column=1, sticky="e");  ttk.Entry(opt, textvariable=self.var_marg_t, width=6).grid(row=4, column=2, sticky="w", padx=6)
        ttk.Label(opt, text="Dolny").grid(row=4, column=3, sticky="e");  ttk.Entry(opt, textvariable=self.var_marg_b, width=6).grid(row=4, column=4, sticky="w", padx=6)

        # --- Sortowanie ---
        sortf = ttk.LabelFrame(parent, text="Sortowanie")
        sortf.pack(fill="x", padx=8, pady=(10,8))
        self.var_sort_col = tk.StringVar(value="")        # pusta = bez sortowania
        self.var_sort_asc = tk.BooleanVar(value=True)     # True=rosnąco, False=malejąco
        ttk.Label(sortf, text="Kolumna:").grid(row=0, column=0, sticky="e", padx=8, pady=6)
        self.cmb_sort_col = ttk.Combobox(sortf, textvariable=self.var_sort_col, width=40, values=[], state="readonly")
        self.cmb_sort_col.grid(row=0, column=1, sticky="w", padx=4, pady=6)
        ttk.Radiobutton(sortf, text="Rosnąco (A→Z / 0→9)", variable=self.var_sort_asc, value=True)\
            .grid(row=0, column=2, sticky="w", padx=12)
        ttk.Radiobutton(sortf, text="Malejąco (Z→A / 9→0)", variable=self.var_sort_asc, value=False)\
            .grid(row=0, column=3, sticky="w", padx=12)

        # --- Układ dokumentu ---
        lay = ttk.LabelFrame(parent, text="Układ dokumentu")
        lay.pack(fill="x", padx=8, pady=(10,8))
        self.var_layout = tk.StringVar(value="A")
        ttk.Radiobutton(lay, text="A — tabela, pod spodem: zdjęcie i/lub mapa", variable=self.var_layout, value="A")\
            .grid(row=0, column=0, sticky="w", padx=8, pady=6)
        ttk.Radiobutton(lay, text="B — tabela po lewej + kolumna na zdjęcie po prawej; mapa pod spodem", variable=self.var_layout, value="B")\
            .grid(row=1, column=0, sticky="w", padx=8, pady=6)

    # ---- Pomocnicze ----
    def _placeholder_in_extra(self, text):
        for w in self.extra_box.winfo_children(): w.destroy()
        ttk.Label(self.extra_box, text=text, foreground="#777", wraplength=900, justify="left")\
            .pack(fill="x", padx=12, pady=12)

    def rebuild_extra_columns(self):
        """Zbuduj listę checkboxów 'Pozostałe kolumny' po wczytaniu presetu."""
        for w in self.extra_box.winfo_children(): w.destroy()
        self.extra_vars = {}
        if self.df is None:
            self._placeholder_in_extra("Wybierz plik Excel, aby rozpocząć."); return
        if not self.mapping:
            self._placeholder_in_extra("Wczytaj preset, aby zobaczyć kolumny pozostające poza mapą."); return
        mapped_sources = {m["source"] for m in self.mapping if m.get("source")}
        any_added = False
        for c in self.df.columns:
            if c in mapped_sources: continue
            v = tk.BooleanVar(value=False)
            ttk.Checkbutton(self.extra_box, text=c, variable=v).pack(anchor="w", padx=8, pady=2)
            self.extra_vars[c] = v; any_added = True
        if not any_added:
            self._placeholder_in_extra("Brak dodatkowych kolumn — preset obejmuje wszystkie źródła.")

    # ---- Presety (GUI) ----
    def load_preset_dialog(self):
        if self.df is None:
            messagebox.showwarning("Najpierw Excel", "Wczytaj plik Excel (żeby dopasować kolumny).")
            return
        p = filedialog.askopenfilename(initialdir="presets",
                                        filetypes=[("Preset JSON","*.json"),("All","*.*")])
        if not p: return
        try:
            data = json.loads(Path(p).read_text(encoding="utf-8"))
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie można wczytać presetu:\n{e}"); return
        aliases = data.get("aliases", {})
        new_mapping = []
        for item in data.get("mapping", []):
            m = dict(item)
            m["source"] = resolve_source_name(m.get("source",""), self.df.columns, aliases)
            new_mapping.append(m)
        self.mapping = new_mapping
        # opcje z presetu (jeśli są)
        opts = data.get("options", {})
        if opts:
            self.var_comma.set(bool(opts.get("decimal_comma", self.var_comma.get())))
            self.var_break.set(bool(opts.get("page_break", self.var_break.get())))
            ph = opts.get("photo", {})
            if "enabled" in ph: self.var_photo.set(bool(ph["enabled"]))
            if "height_cm" in ph: self.var_photo_h.set(float(ph["height_cm"]))
            mp = opts.get("map", {})
            if "enabled" in mp: self.var_map.set(bool(mp["enabled"]))
            if "height_cm" in mp: self.var_map_h.set(float(mp["height_cm"]))
        # UI
        self.preset_label_var.set(data.get("name","(preset)"))
        self.refresh_tree()
        self.rebuild_extra_columns()
        self.cmb_source.configure(values=[""] + list(self.df.columns))

    # ---- Handlery standardowe ----
    def pick_excel(self):
        p = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx *.xls"), ("All","*.*")])
        if not p: return
        self.in_entry.delete(0, tk.END); self.in_entry.insert(0, p)
        # Automatyczne wczytanie kolumn po wyborze pliku
        self.load_columns()

    def pick_outdir(self):
        d = filedialog.askdirectory()
        if d: self.out_entry.delete(0, tk.END); self.out_entry.insert(0, d)

    def load_columns(self):
        path = self.in_entry.get().strip()
        if not path:
            messagebox.showwarning("Brak pliku","Wskaż plik Excel."); return
        try:
            self.df = read_excel_any(path)
        except Exception as e:
            messagebox.showerror("Błąd wczytywania", f"Nie udało się wczytać Excela:\n{e}\n\nJeśli to .xls, wymagany xlrd>=2.0.1.")
            return
        # start bez mapy (czysta lista) – preset wczytujesz ręcznie
        self.mapping = []
        self.preset_label_var.set("(brak)")
        self.refresh_tree()
        # „Pozostałe kolumny” – placeholder do czasu wczytania presetu
        self.rebuild_extra_columns()
        # comboboxy z listą kolumn
        cols_list = [""] + list(self.df.columns)
        self.cmb_source.configure(values=cols_list)
        # sort: kolumna i reset kierunku
        self.cmb_sort_col.configure(values=cols_list)
        self.var_sort_col.set("")

    def refresh_tree(self):
        self.tree.delete(*self.tree.get_children())
        for i, m in enumerate(self.mapping):
            self.tree.insert("", "end", iid=str(i),
                values=(i, "✓" if m.get("enabled",False) else "–",
                        m.get("label",""), m.get("source",""),
                        m.get("transform","identity"), m.get("const","")))

    def on_select_row(self, *_):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0]); m = self.mapping[i]
        self.var_on.set(bool(m.get("enabled",False)))
        self.var_label.set(m.get("label","")); self.var_source.set(m.get("source",""))
        self.var_transform.set(m.get("transform","identity")); self.var_const.set(m.get("const",""))

    def apply_edit(self):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0])
        self.mapping[i] = {
            "enabled": bool(self.var_on.get()),
            "label": self.var_label.get().strip(),
            "source": self.var_source.get().strip(),
            "transform": self.var_transform.get().strip(),
            "const": self.var_const.get(),
        }
        self.refresh_tree(); self.tree.selection_set(str(i))
        self.rebuild_extra_columns()

    def move_selected(self, delta):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0]); j = i + delta
        if j < 0 or j >= len(self.mapping): return
        self.mapping[i], self.mapping[j] = self.mapping[j], self.mapping[i]
        self.refresh_tree(); self.tree.selection_set(str(j))

    def delete_selected(self):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0]); del self.mapping[i]; self.refresh_tree()
        self.rebuild_extra_columns()

    def add_from_column(self):
        src = self.var_source.get().strip()
        if not src:
            messagebox.showinfo("Info","Wybierz źródło w polu 'Źródło', potem kliknij 'Dodaj z kolumny…'."); return
        self.mapping.append({"enabled": False, "label": src, "source": src, "transform":"identity", "const":""})
        self.refresh_tree(); self.tree.selection_set(str(len(self.mapping)-1))
        self.rebuild_extra_columns()

    def add_blank_row(self):
        self.mapping.append({"enabled": True, "label": "", "source": "", "transform":"constant", "const":""})
        self.refresh_tree(); self.tree.selection_set(str(len(self.mapping)-1))

    def run(self):
        if self.df is None:
            messagebox.showwarning("Brak danych","Najpierw wybierz plik Excel."); return
        out_dir = self.out_entry.get().strip() or os.path.dirname(self.in_entry.get()) or os.getcwd()
        if not os.path.isdir(out_dir):
            messagebox.showwarning("Folder","Podany folder wyjściowy nie istnieje."); return
        out_name = self.name_entry.get().strip() or "wynik.docx"
        extra = [k for k,v in self.extra_vars.items() if v.get()]
        if not self.mapping and not extra:
            messagebox.showwarning("Brak pól", "Mapa jest pusta i nie wybrano żadnych dodatkowych kolumn.\n"
                                               "Wczytaj preset lub dodaj wiersze/kolumny.")
            return

        # Sortowanie (opcjonalne)
        df_use = self.df
        col = self.var_sort_col.get().strip()
        if col:
            asc = self.var_sort_asc.get()
            def _sort_key(s):
                # Jeśli kolumna jest tekstowa, spróbuj sortować liczbowo (z przecinkami/kropkami)
                if s.dtype == "O":
                    s_num = pd.to_numeric(s.astype(str).str.replace(",", ".", regex=False), errors="coerce")
                    if s_num.notna().any():
                        return s_num
                return s
            try:
                df_use = self.df.sort_values(by=col, ascending=asc, na_position="last", key=_sort_key)
            except Exception:
                df_use = self.df

        try:
            out_file = build_docx(
                df_use,
                list(self.mapping), extra, out_dir, out_name,
                add_photo=self.var_photo.get(),
                photo_h_cm=float(self.var_photo_h.get()),
                add_map=self.var_map.get(),
                map_h_cm=float(self.var_map_h.get()),
                page_break=self.var_break.get(),
                use_decimal_comma=self.var_comma.get(),
                margin_left_cm=float(self.var_marg_l.get()),
                margin_right_cm=float(self.var_marg_r.get()),
                margin_top_cm=float(self.var_marg_t.get()),
                margin_bottom_cm=float(self.var_marg_b.get()),
                layout_mode=self.var_layout.get(),
            )
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się wygenerować DOCX:\n{e}"); return
        messagebox.showinfo("Gotowe", f"Zapisano:\n{out_file}")

if __name__ == "__main__":
    App().mainloop()