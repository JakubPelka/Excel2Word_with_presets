# -*- coding: utf-8 -*-
"""
Excel → mikro-tabele w Wordzie (GUI, Tkinter) z presetami, układem A/B,
auto-zdjęciami (match po kolumnie) + szerokościami kolumn + czcionką + LOG + KOMPRESJA
+ podział ramki zdjęcia na 2 pola
+ foto #2: sufiksy _2, -2, (2),  (2) oraz centrowanie zdjęć

DOMYŚLNE:
- JPG=90, DPI=450
- Arial 9 pt
- Marginesy: L=2.0, P=6.5, G=3.9, D=3.0

NOWOŚĆ:
- pierwszy wiersz każdej głównej tabeli = nagłówek: tło #A6A6A6, biały pogrubiony tekst
- transform "format": łączenie wartości z wielu kolumn wg szablonu (const) i "sources"
"""

# ---------- BOOTSTRAP PIP ----------
import sys, subprocess
def _pip_install(spec):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-U", spec])
    except Exception:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "-U", spec])

def _ensure_min(pkg, module, min_v):
    try:
        try:
            from importlib.metadata import version as _v
        except Exception:
            from importlib_metadata import version as _v
        cur = _v(module)
    except Exception:
        cur = None
    def _lt(a,b):
        pa=lambda s:[int(x) for x in str(s).split(".") if x.isdigit()]
        return pa(a or "0")<pa(b)
    if cur is None or _lt(cur, min_v):
        _pip_install(f"{pkg}>={min_v}")

def _ensure(pkg, module=None):
    import importlib.util
    if not importlib.util.find_spec(module or pkg.replace("-","_")):
        _pip_install(pkg)

_ensure("pandas","pandas"); _ensure("python-docx","docx"); _ensure("openpyxl","openpyxl"); _ensure_min("xlrd","xlrd","2.0.1")
_ensure("Pillow","PIL")

# ---------- IMPORTY ----------
import os, json, logging, traceback, io
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date, datetime
from PIL import Image, ImageOps

EMU_PER_CM = 360000  # Word EMU -> cm

# ---------- LOGGING ----------
def setup_logger(out_dir: str):
    out_dir = out_dir or os.getcwd()
    try:
        os.makedirs(out_dir, exist_ok=True)
    except Exception:
        pass
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_path = os.path.join(out_dir, f"excel2word_{ts}.log")
    logger = logging.getLogger(f"e2w_{ts}")
    logger.setLevel(logging.DEBUG)
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    sh = logging.StreamHandler()
    sh.setLevel(logging.WARNING)
    sh.setFormatter(fmt)
    logger.addHandler(sh)
    logger.log_path = log_path
    logger.debug("=== START SESSION ===")
    return logger

# ---------- TRANSFORMACJE ----------
def _is_missing(v):
    return v is None or (isinstance(v, float) and pd.isna(v)) or (isinstance(v, str) and v.strip()=="")

def tf_identity(val, row, comma=False):
    return "" if _is_missing(val) else str(val)

def tf_m2_to_ha_round2(val, row, comma=False):
    if _is_missing(val): return ""
    try:
        ha = round(float(str(val).replace(",", "."))/10000.0, 2)
        s = f"{ha:.2f}"
        return s.replace(".", ",") if comma else s
    except Exception:
        return ""

def tf_prelim_to_bedomning(val, row, comma=False):
    v = "" if _is_missing(val) else str(val).strip().lower()
    return "Säker" if v in {"0","nej","no","false","0.0",""} else "Preliminärt"

def tf_constant(val, row, comma=False, const_val=""):
    return const_val

def tf_date_only(val, row, comma=False):
    if _is_missing(val): return ""
    try:
        if isinstance(val, (datetime, date, pd.Timestamp)):
            return pd.to_datetime(val).date().isoformat()
        s = str(val).strip()
        if " " in s: s = s.split(" ")[0]
        if "T" in s: s = s.split("T")[0]
        d = pd.to_datetime(s, errors="coerce")
        return d.date().isoformat() if not pd.isna(d) else s
    except Exception:
        return str(val)

def tf_bool_ja_nej(val, row, comma=False):
    if _is_missing(val): return ""
    s = str(val).strip().lower()
    try:
        f = float(s.replace(",", "."))
        return "Ja" if f != 0.0 else "Nej"
    except Exception:
        pass
    positives = {"ja","yes","y","true","t","1"}
    negatives = {"nej","no","n","false","f","0"}
    if s in positives: return "Ja"
    if s in negatives: return "Nej"
    return ""

def tf_format(val, row, comma=False, const_val="", sources=None):
    """
    Składa tekst wg szablonu const_val z wielu kolumn (sources).
    Przykład:
      sources = ["xcoord", "ycoord"]
      const_val = "{xcoord}, {ycoord}"
    """
    sources = sources or []
    data = {}
    for name in sources:
        v = row.get(name, "")
        if v is None or (isinstance(v, float) and pd.isna(v)):
            v = ""
        else:
            v = str(v)
            if comma:
                # spróbuj zamienić . na , dla liczb
                try:
                    float(v.replace(",", "."))
                    v = v.replace(".", ",")
                except Exception:
                    pass
        data[name] = v
    fmt = const_val or " ".join("{" + s + "}" for s in sources)
    try:
        out = fmt.format(**data)
    except Exception:
        out = " ".join(data.get(s, "") for s in sources)
    return out.strip()

TRANSFORMS = {
    "identity": ("Bez zmian", tf_identity),
    "m2_to_ha_round2": ("m² → ha (0,01)", tf_m2_to_ha_round2),
    "prelim_to_bedomning": ("0/nej → Säker, inne → Preliminärt", tf_prelim_to_bedomning),
    "constant": ("Stała wartość", tf_constant),
    "date_only": ("Tylko data (YYYY-MM-DD)", tf_date_only),
    "bool_ja_nej": ("Bool → Ja/Nej", tf_bool_ja_nej),
    "format": ("Formatuj z wielu kolumn", tf_format),
}

# ---------- DOCX HELPERS ----------
def _shade_cell(cell, fill_hex="F2F2F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def _set_cell_text(cell, text, bold=False, size_pt=None, font_name=None, color_rgb=None):
    cell.text = "" if text is None else str(text)
    for p in cell.paragraphs:
        for r in p.runs:
            if font_name: r.font.name = font_name
            if size_pt:   r.font.size = Pt(size_pt)
            if color_rgb: r.font.color.rgb = RGBColor(*color_rgb)
            r.font.bold = bold
    if cell.paragraphs:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

def _set_tbl_borders(table, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4, color="000000"):
    tbl = table._element  # CT_Tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.append(tblPr)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders'); tblPr.append(tblBorders)
    def set_edge(tag, val):
        e = tblBorders.find(qn(f"w:{tag}"))
        if e is None:
            e = OxmlElement(f"w:{tag}"); tblBorders.append(e)
        if val == 0:
            e.set(qn("w:val"), "nil")
        else:
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(val))
            e.set(qn("w:color"), color); e.set(qn("w:space"), "0")
    set_edge("top", top); set_edge("left", left); set_edge("bottom", bottom); set_edge("right", right)
    set_edge("insideH", insideH); set_edge("insideV", insideV)

def _style_header_row(table, font_name, font_size_pt):
    """Nadaje styl nagłówkowy pierwszemu wierszowi: #A6A6A6 + biały pogrubiony tekst."""
    if not table.rows: return
    row0 = table.rows[0]
    for c in row0.cells:
        _shade_cell(c, "A6A6A6")
        _set_cell_text(c, c.text, bold=True, size_pt=font_size_pt, font_name=font_name, color_rgb=(0xFF,0xFF,0xFF))

# ---------- OBRAZY ----------
def _normalize_slug_candidates(val):
    cands = []
    if val is None: return cands
    s = str(val).strip()
    if s == "": return cands
    cands.append(s)
    cands.append(s.replace(" ", "_"))
    try:
        f = float(s.replace(",", "."))
        if f.is_integer():
            cands.append(str(int(f)))
    except Exception:
        pass
    cands.append(s.strip(" ._"))
    dedup = []
    for x in cands:
        if x and x not in dedup:
            dedup.append(x)
    return dedup

def find_image_for(id_value, images_dir, logger=None):
    """Foto #1: bazowy plik odpowiadający ID."""
    if logger: logger.debug(f"[IMG] images_dir={images_dir!r}, id_value={id_value!r}")
    if not images_dir or not os.path.isdir(images_dir) or _is_missing(id_value):
        if logger: logger.debug("[IMG] brak folderu lub id")
        return None
    cands = _normalize_slug_candidates(id_value)
    if logger: logger.debug(f"[IMG] candidates={cands}")
    for base in cands:
        for ext in ("jpg","jpeg","png","JPG","JPEG","PNG"):
            p = os.path.join(images_dir, f"{base}.{ext}")
            if os.path.exists(p):
                if logger: logger.debug(f"[IMG] FOUND {p}")
                return p
    try:
        slug_set = {x.lower() for x in cands}
        for fn in os.listdir(images_dir):
            base, ext = os.path.splitext(fn)
            if ext.lower() in (".jpg",".jpeg",".png") and base.lower() in slug_set:
                p = os.path.join(images_dir, fn)
                if logger: logger.debug(f"[IMG] FOUND (scan) {p}")
                return p
    except Exception as e:
        if logger: logger.warning(f"[IMG] scan error: {e}")
    if logger: logger.debug("[IMG] NOT FOUND")
    return None

def find_image_for_second(id_value, images_dir, logger=None, suffixes=None):
    """
    Foto #2: warianty <base><suffix>.*  Domyślne sufiksy: "_2", "-2", "(2)", " (2)".
    """
    if suffixes is None:
        suffixes = ("_2", "-2", "(2)", " (2)")
    if logger: logger.debug(f"[IMG2] images_dir={images_dir!r}, id_value={id_value!r}, suffixes={suffixes}")
    if not images_dir or not os.path.isdir(images_dir) or _is_missing(id_value):
        if logger: logger.debug("[IMG2] brak folderu lub id")
        return None
    bases = _normalize_slug_candidates(id_value)  # np. ["Lokal 2","Lokal_2","2"]
    cands = []
    for b in bases:
        for sfx in suffixes:
            cands.append(b + sfx)
    if logger: logger.debug(f"[IMG2] candidates={cands}")
    for base in cands:
        for ext in ("jpg","jpeg","png","JPG","JPEG","PNG"):
            p = os.path.join(images_dir, f"{base}.{ext}")
            if os.path.exists(p):
                if logger: logger.debug(f"[IMG2] FOUND {p}")
                return p
    try:
        want = {x.lower() for x in cands}
        for fn in os.listdir(images_dir):
            base, ext = os.path.splitext(fn)
            if ext.lower() in (".jpg",".jpeg",".png") and base.lower() in want:
                p = os.path.join(images_dir, fn)
                if logger: logger.debug(f"[IMG2] FOUND (scan) {p}")
                return p
    except Exception as e:
        if logger: logger.warning(f"[IMG2] scan error: {e}")
    if logger: logger.debug("[IMG2] NOT FOUND")
    return None

def _cm_to_px(cm, dpi):
    return int(round(cm * dpi / 2.54))

def insert_picture_in_cell(cell, image_path, frame_w_cm, frame_h_cm=None, logger=None,
                           jpg_quality=90, export_dpi=450, no_upscale=True):
    """Skaluje obraz do ramki, re-koduje do JPEG (bez EXIF), wstawia i centruje poziomo."""
    if not image_path or not os.path.exists(image_path):
        if logger: logger.debug(f"[IMG] no image to insert: {image_path}")
        return False
    try:
        with Image.open(image_path) as im:
            im = ImageOps.exif_transpose(im).convert("RGB")
            src_w, src_h = im.size

            tgt_w_px = _cm_to_px(frame_w_cm, export_dpi)
            tgt_h_px = _cm_to_px(frame_h_cm, export_dpi) if frame_h_cm else 10**9

            scale = min(tgt_w_px / src_w, tgt_h_px / src_h)
            if no_upscale:
                scale = min(1.0, scale)
            new_w = max(1, int(round(src_w * scale)))
            new_h = max(1, int(round(src_h * scale)))

            if (new_w, new_h) != (src_w, src_h):
                im = im.resize((new_w, new_h), Image.LANCZOS)

            bio = io.BytesIO()
            im.save(bio, format="JPEG", quality=int(jpg_quality), optimize=True,
                    progressive=True, subsampling=2)
            bio.seek(0)

            new_w_cm = new_w / export_dpi * 2.54
            target_w_cm = min(frame_w_cm, new_w_cm)

        try:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            pass
        p = cell.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(bio, width=Cm(max(0.1, target_w_cm)))
        return True
    except Exception as e:
        if logger:
            logger.error(f"[IMG] insert failed for {image_path}: {e}")
            logger.debug(traceback.format_exc())
        return False

# ---------- EXCEL ----------
def read_excel_any(path):
    ext = os.path.splitext(path.lower())[1]
    if ext == ".xls": return pd.read_excel(path, engine="xlrd")
    return pd.read_excel(path)

# ---------- PRESETY / ALIASY ----------
def _norm(s: str) -> str:
    return "".join(ch for ch in str(s).lower() if ch.isalnum())

def resolve_source_name(wanted: str, df_columns, aliases: dict):
    cols = list(map(str, df_columns))
    by_norm = {_norm(c): c for c in cols}
    candidates = [wanted] + list(aliases.get(wanted, [])) if wanted else []
    for cand in candidates:
        n = _norm(cand)
        if n in by_norm: return by_norm[n]
        for c in cols:
            if n and n in _norm(c): return c
    return ""

# ---------- GENERACJA DOCX ----------
def build_docx(df, mapping_rows, extra_cols, out_path, out_name,
               add_photo=True, photo_h_cm=6.0, photo_split=False,
               add_map=True, map_h_cm=6.0,
               page_break=False, use_decimal_comma=True,
               margin_left_cm=2.0, margin_right_cm=6.5,
               margin_top_cm=3.9, margin_bottom_cm=3.0,
               layout_mode="A",
               images_dir=None, image_id_column="objektnummer",
               base_font_name="Arial", base_font_size_pt=9,
               a_label_cm=6.0,
               b_label_cm=0.0, b_value_cm=0.0, b_photo_cm=0.0,
               jpg_quality=90, export_dpi=450, no_upscale=True,
               logger=None):
    if logger:
        logger.debug(f"[DOCX] start build_docx out_name={out_name}, layout={layout_mode}, "
                     f"photo={add_photo}/{photo_h_cm}cm split={photo_split}, map={add_map}/{map_h_cm}cm, "
                     f"page_break={page_break}, comma={use_decimal_comma}, "
                     f"margins=({margin_left_cm},{margin_right_cm},{margin_top_cm},{margin_bottom_cm})cm, "
                     f"font=({base_font_name},{base_font_size_pt}pt), "
                     f"A_label={a_label_cm}cm, B={b_label_cm},{b_value_cm},{b_photo_cm}cm, "
                     f"images_dir={images_dir}, image_id_column={image_id_column}, "
                     f"jpg_quality={jpg_quality}, export_dpi={export_dpi}, no_upscale={no_upscale}")
    doc = Document()
    try:
        style = doc.styles["Normal"]
        if base_font_name: style.font.name = base_font_name
        if base_font_size_pt: style.font.size = Pt(base_font_size_pt)
    except Exception as e:
        if logger: logger.warning(f"[DOCX] style set failed: {e}")

    section = doc.sections[0]
    section.left_margin   = Cm(float(margin_left_cm))
    section.right_margin  = Cm(float(margin_right_cm))
    section.top_margin    = Cm(float(margin_top_cm))
    section.bottom_margin = Cm(float(margin_bottom_cm))

    content_w_cm = (section.page_width - section.left_margin - section.right_margin) / EMU_PER_CM

    # Dołącz pozostałe kolumny jako identity
    for col_name in extra_cols:
        mapping_rows.append({"enabled": True, "label": col_name, "source": col_name, "transform": "identity", "const": ""})

    for ridx, (_, cur_row) in enumerate(df.iterrows(), start=1):
        if logger: logger.debug(f"[ROW] #{ridx}")
        enabled_rows = [m for m in mapping_rows if m.get("enabled", False)]
        img_path = None; img2_path = None
        if images_dir and image_id_column:
            id_val = cur_row.get(image_id_column, None)
            img_path = find_image_for(id_val, images_dir, logger=logger)
            img2_path = find_image_for_second(id_val, images_dir, logger=logger)

        def _compute_value(item):
            lbl = item.get("label",""); src = item.get("source","")
            tf_key = item.get("transform","identity"); const_val = item.get("const","")
            sources = item.get("sources", [])
            raw = cur_row.get(src, "") if src else ""
            tf = TRANSFORMS.get(tf_key, TRANSFORMS["identity"])[1]
            try:
                val = tf(raw, cur_row, use_decimal_comma, const_val=const_val, sources=sources)
            except TypeError:
                # kompatybilnie z transformami nieobsługującymi 'sources' / 'const_val'
                if tf_key == "constant":
                    val = tf(raw, cur_row, use_decimal_comma, const_val=const_val)
                else:
                    val = tf(raw, cur_row, use_decimal_comma)
            if tf_key != "constant" and (val is None or (isinstance(val,str) and val.strip()=="")):
                val = " - "
            return lbl, val

        if layout_mode == "B" and add_photo and enabled_rows:
            # szerokości B (0=auto)
            photo_cm = b_photo_cm if (b_photo_cm and b_photo_cm > 0) else max(5.0, min(9.0, content_w_cm * 0.35))
            left_cm  = max(6.0, content_w_cm - photo_cm)
            label_cm = b_label_cm if (b_label_cm and b_label_cm > 0) else max(3.5, min(7.0, left_cm * 0.45))
            value_cm = b_value_cm if (b_value_cm and b_value_cm > 0) else max(3.5, left_cm - label_cm)
            if b_label_cm > 0 and b_value_cm > 0 and (b_label_cm + b_value_cm) > left_cm:
                value_cm = max(1.0, left_cm - b_label_cm)

            t = doc.add_table(rows=len(enabled_rows), cols=3)
            t.style = "Table Grid"; t.autofit = False; t.alignment = WD_TABLE_ALIGNMENT.LEFT

            for i, item in enumerate(enabled_rows):
                lbl, val = _compute_value(item)
                _set_cell_text(t.cell(i,0), lbl, bold=True, size_pt=base_font_size_pt, font_name=base_font_name); _shade_cell(t.cell(i,0), "F2F2F2")
                _set_cell_text(t.cell(i,1), val, bold=False, size_pt=base_font_size_pt, font_name=base_font_name)

            for r in t.rows:
                r.cells[0].width = Cm(label_cm); r.cells[1].width = Cm(value_cm); r.cells[2].width = Cm(photo_cm)

            # NAGŁÓWEK — CAŁY pierwszy wiersz:
            _style_header_row(t, base_font_name, base_font_size_pt)

            # Kolumna zdjęcia (scalona) + podpis
            merged = t.cell(0,2)
            for i in range(1, len(enabled_rows)):
                merged = merged.merge(t.cell(i,2))
            p = merged.paragraphs[0] if merged.paragraphs else merged.add_paragraph("")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("Representativt foto:")
            if base_font_name: run.font.name = base_font_name
            if base_font_size_pt: run.font.size = Pt(base_font_size_pt)

            # zdjęcia
            if photo_split:
                inner = merged.add_table(rows=1, cols=2)
                inner.style = "Table Grid"; inner.autofit = False
                inner.rows[0].height = Cm(photo_h_cm)
                inner.cell(0,0).width = Cm(photo_cm/2.0)
                inner.cell(0,1).width = Cm(photo_cm/2.0)
                if img_path:
                    insert_picture_in_cell(inner.cell(0,0), img_path, frame_w_cm=photo_cm/2.0, frame_h_cm=photo_h_cm,
                                           logger=logger, jpg_quality=jpg_quality, export_dpi=export_dpi, no_upscale=no_upscale)
                if img2_path:
                    insert_picture_in_cell(inner.cell(0,1), img2_path, frame_w_cm=photo_cm/2.0, frame_h_cm=photo_h_cm,
                                           logger=logger, jpg_quality=jpg_quality, export_dpi=export_dpi, no_upscale=no_upscale)
                _set_tbl_borders(inner, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)
            else:
                if img_path:
                    insert_picture_in_cell(merged, img_path, frame_w_cm=photo_cm, frame_h_cm=None,
                                           logger=logger, jpg_quality=jpg_quality, export_dpi=export_dpi, no_upscale=no_upscale)

            _set_tbl_borders(t, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)

            doc.add_paragraph("")
            if add_map:
                p2 = doc.add_paragraph("Kartbild av objektet:")
                if p2.runs:
                    if base_font_name: p2.runs[0].font.name = base_font_name
                    if base_font_size_pt: p2.runs[0].font.size = Pt(base_font_size_pt)
                mp = doc.add_table(rows=1, cols=1); mp.style = "Table Grid"; mp.autofit = False
                mp.rows[0].height = Cm(map_h_cm); mp.cell(0, 0).width = Cm(content_w_cm)
                _set_cell_text(mp.cell(0, 0), " ", size_pt=base_font_size_pt, font_name=base_font_name)
                _set_tbl_borders(mp, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)
                doc.add_paragraph("")

        else:
            # Układ A
            label_cm = max(1.0, min(content_w_cm-1.0, a_label_cm if a_label_cm else 6.0))
            value_cm = max(1.0, content_w_cm - label_cm)

            t = doc.add_table(rows=0, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = "Table Grid"; t.autofit = False
            for item in enabled_rows or mapping_rows:
                if not item.get("enabled", False): continue
                lbl, val = _compute_value(item)
                cells = t.add_row().cells
                _set_cell_text(cells[0], lbl, bold=True, size_pt=base_font_size_pt, font_name=base_font_name); _shade_cell(cells[0], "F2F2F2")
                _set_cell_text(cells[1], val, bold=False, size_pt=base_font_size_pt, font_name=base_font_name)
            for r in t.rows:
                r.cells[0].width = Cm(label_cm); r.cells[1].width = Cm(value_cm)

            # NAGŁÓWEK — CAŁY pierwszy wiersz:
            _style_header_row(t, base_font_name, base_font_size_pt)

            _set_tbl_borders(t, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)

            doc.add_paragraph("")
            if add_photo:
                p = doc.add_paragraph("Representativt foto:")
                if p.runs:
                    if base_font_name: p.runs[0].font.name = base_font_name
                    if base_font_size_pt: p.runs[0].font.size = Pt(base_font_size_pt)
                if photo_split:
                    ph = doc.add_table(rows=1, cols=2); ph.style = "Table Grid"; ph.autofit = False
                    ph.rows[0].height = Cm(photo_h_cm)
                    ph.cell(0, 0).width = Cm(content_w_cm/2.0)
                    ph.cell(0, 1).width = Cm(content_w_cm/2.0)
                    if img_path:
                        insert_picture_in_cell(ph.cell(0,0), img_path, frame_w_cm=content_w_cm/2.0, frame_h_cm=photo_h_cm,
                                               logger=logger, jpg_quality=jpg_quality, export_dpi=export_dpi, no_upscale=no_upscale)
                    if img2_path:
                        insert_picture_in_cell(ph.cell(0,1), img2_path, frame_w_cm=content_w_cm/2.0, frame_h_cm=photo_h_cm,
                                               logger=logger, jpg_quality=jpg_quality, export_dpi=export_dpi, no_upscale=no_upscale)
                    _set_tbl_borders(ph, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)
                else:
                    ph = doc.add_table(rows=1, cols=1); ph.style = "Table Grid"; ph.autofit = False
                    ph.rows[0].height = Cm(photo_h_cm); ph.cell(0, 0).width = Cm(content_w_cm)
                    if img_path:
                        insert_picture_in_cell(ph.cell(0,0), img_path, frame_w_cm=content_w_cm, frame_h_cm=photo_h_cm,
                                               logger=logger, jpg_quality=jpg_quality, export_dpi=export_dpi, no_upscale=no_upscale)
                    else:
                        _set_cell_text(ph.cell(0, 0), " ", size_pt=base_font_size_pt, font_name=base_font_name)
                    _set_tbl_borders(ph, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)
                doc.add_paragraph("")
            if add_map:
                p2 = doc.add_paragraph("Kartbild av objektet:")
                if p2.runs:
                    if base_font_name: p2.runs[0].font.name = base_font_name
                    if base_font_size_pt: p2.runs[0].font.size = Pt(base_font_size_pt)
                mp = doc.add_table(rows=1, cols=1); mp.style = "Table Grid"; mp.autofit = False
                mp.rows[0].height = Cm(map_h_cm); mp.cell(0, 0).width = Cm(content_w_cm)
                _set_cell_text(mp.cell(0, 0), " ", size_pt=base_font_size_pt, font_name=base_font_name)
                _set_tbl_borders(mp, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)
                doc.add_paragraph("")

        if page_break:
            doc.add_page_break()

    out_file = os.path.join(out_path, out_name if out_name.lower().endswith(".docx") else out_name + ".docx")
    doc.save(out_file)
    if logger: logger.debug(f"[DOCX] saved: {out_file}")
    return out_file

# ---------- GUI ----------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Excel → Mikro-tabele Word — presety, układ A/B, 2-zdjęcia (_2/-2/(2)), foto-match, kolumny, czcionka, log, kompresja, format()")
        self.geometry("1180x980"); self.resizable(True, True)
        self.df = None
        self.mapping = []
        self.extra_vars = {}
        self.preset_label_var = tk.StringVar(value="(brak)")
        # zdjęcia
        self.var_img_dir = tk.StringVar(value="")
        self.var_img_col = tk.StringVar(value="objektnummer")
        # foto-kompresja — DOMYŚLNE
        self.var_jpg_quality = tk.IntVar(value=90)
        self.var_export_dpi = tk.IntVar(value=450)
        self.var_no_upscale = tk.BooleanVar(value=True)
        # foto: podział na 2 pola
        self.var_photo_split = tk.BooleanVar(value=True)
        # czcionka — DOMYŚLNE
        self.var_font_name = tk.StringVar(value="Arial")
        self.var_font_size = tk.DoubleVar(value=9.0)
        # szerokości kolumn
        self.var_A_label_cm = tk.DoubleVar(value=6.0)
        self.var_B_label_cm = tk.DoubleVar(value=0.0)   # 0 = auto
        self.var_B_value_cm = tk.DoubleVar(value=0.0)   # 0 = auto
        self.var_B_photo_cm = tk.DoubleVar(value=0.0)   # 0 = auto
        self._build_ui()

    def _build_ui(self):
        pad = {"padx":8, "pady":6}
        top = ttk.Frame(self); top.pack(fill="both", expand=True)
        ttk.Label(top, text="Plik Excel (.xlsx / .xls):").grid(row=0, column=0, sticky="w", **pad)
        self.in_entry = ttk.Entry(top, width=80); self.in_entry.grid(row=0, column=1, sticky="we", **pad)
        ttk.Button(top, text="Wybierz…", command=self.pick_excel).grid(row=0, column=2, **pad)
        ttk.Label(top, text="Folder wyjściowy:").grid(row=1, column=0, sticky="w", **pad)
        self.out_entry = ttk.Entry(top, width=80); self.out_entry.grid(row=1, column=1, sticky="we", **pad)
        ttk.Button(top, text="Wybierz…", command=self.pick_outdir).grid(row=1, column=2, **pad)
        ttk.Label(top, text="Nazwa pliku .docx:").grid(row=2, column=0, sticky="w", **pad)
        self.name_entry = ttk.Entry(top, width=40); self.name_entry.insert(0, "wynik.docx"); self.name_entry.grid(row=2, column=1, sticky="w", **pad)

        # Wiersz: zdjęcia
        ttk.Label(top, text="Folder zdjęć (opcjonalnie):").grid(row=3, column=0, sticky="w", **pad)
        self.img_entry = ttk.Entry(top, textvariable=self.var_img_dir, width=80); self.img_entry.grid(row=3, column=1, sticky="we", **pad)
        ttk.Button(top, text="Wybierz…", command=self.pick_imgdir).grid(row=3, column=2, **pad)

        ttk.Label(top, text="Kolumna dopasowania zdjęć:").grid(row=4, column=0, sticky="w", **pad)
        self.cmb_img_col = ttk.Combobox(top, textvariable=self.var_img_col, width=40, values=[], state="readonly")
        self.cmb_img_col.grid(row=4, column=1, sticky="w", **pad)

        btn_row = ttk.Frame(top); btn_row.grid(row=5, column=0, columnspan=3, sticky="we", padx=8, pady=(2,8))
        ttk.Button(btn_row, text="Wczytaj preset…", command=self.load_preset_dialog).pack(side="left")
        ttk.Label(btn_row, text="Preset:").pack(side="left", padx=(16,4))
        ttk.Label(btn_row, textvariable=self.preset_label_var, foreground="#555").pack(side="left")

        nb = ttk.Notebook(top); nb.grid(row=6, column=0, columnspan=3, sticky="nsew", padx=8, pady=8)
        top.rowconfigure(6, weight=1); top.columnconfigure(1, weight=1)

        # --- Mapowanie
        self.tab_map = ttk.Frame(nb); nb.add(self.tab_map, text="Mapowanie pól")
        self._build_map_tab(self.tab_map)

        # --- Pozostałe kolumny
        self.tab_extra = ttk.Frame(nb); nb.add(self.tab_extra, text="Pozostałe kolumny")
        self.extra_box = ttk.LabelFrame(self.tab_extra, text="Kolumny (nieużyte w mapie)")
        self.extra_box.pack(fill="both", expand=True, padx=8, pady=8)

        # --- Opcje
        self.tab_opt = ttk.Frame(nb); nb.add(self.tab_opt, text="Opcje")
        self._build_opt_tab(self.tab_opt)

        ttk.Button(top, text="Generuj DOCX", command=self.run)\
            .grid(row=7, column=0, columnspan=3, sticky="we", padx=8, pady=8)

    def _build_map_tab(self, parent):
        cols = ("order","#on","label","source","transform","const")
        self.tree = ttk.Treeview(parent, columns=cols, show="headings", height=12)
        for c, txt, w in [
            ("order","Kolej.",60),
            ("#on","Włącz",60),
            ("label","Nowa nazwa",260),
            ("source","Źródło",320),
            ("transform","Transformacja",200),
            ("const","Stała / Format",260),
        ]:
            self.tree.heading(c, text=txt); self.tree.column(c, width=w, anchor="w")
        self.tree.pack(fill="both", expand=True, padx=8, pady=8)

        ed = ttk.LabelFrame(parent, text="Edytuj / dodaj wiersz")
        ed.pack(fill="x", padx=8, pady=(0,8))
        self.var_on = tk.BooleanVar(value=True)
        self.var_label = tk.StringVar(); self.var_source = tk.StringVar()
        self.var_transform = tk.StringVar(value="identity"); self.var_const = tk.StringVar()

        ttk.Checkbutton(ed, text="Włączony", variable=self.var_on).grid(row=0, column=0, sticky="w", padx=6, pady=4)
        ttk.Label(ed, text="Nowa nazwa:").grid(row=0, column=1, sticky="e", padx=6)
        self.ent_label = ttk.Entry(ed, textvariable=self.var_label, width=34); self.ent_label.grid(row=0, column=2, sticky="w", padx=6)

        ttk.Label(ed, text="Źródło:").grid(row=1, column=1, sticky="e", padx=6)
        self.cmb_source = ttk.Combobox(ed, textvariable=self.var_source, width=32, values=[], state="readonly")
        self.cmb_source.grid(row=1, column=2, sticky="w", padx=6)

        ttk.Label(ed, text="Transformacja:").grid(row=0, column=3, sticky="e", padx=6)
        self.cmb_transform = ttk.Combobox(ed, textvariable=self.var_transform, width=28,
                                          values=[k for k in TRANSFORMS], state="readonly")
        self.cmb_transform.grid(row=0, column=4, sticky="w", padx=6)

        ttk.Label(ed, text="Stała / Format:").grid(row=1, column=3, sticky="e", padx=6)
        self.ent_const = ttk.Entry(ed, textvariable=self.var_const, width=30); self.ent_const.grid(row=1, column=4, sticky="w", padx=6)

        ttk.Button(ed, text="Zastosuj (aktualizuj wybrany)", command=self.apply_edit)\
            .grid(row=0, column=5, rowspan=2, sticky="nsw", padx=6, pady=4)

        btns = ttk.Frame(parent); btns.pack(fill="x", padx=8, pady=(0,4))
        ttk.Button(btns, text="↑", width=3, command=lambda: self.move_selected(-1)).pack(side="left", padx=2)
        ttk.Button(btns, text="↓", width=3, command=lambda: self.move_selected(1)).pack(side="left", padx=2)
        ttk.Button(btns, text="Usuń", command=self.delete_selected).pack(side="left", padx=8)
        ttk.Button(btns, text="Dodaj z kolumny…", command=self.add_from_column).pack(side="left", padx=8)
        ttk.Button(btns, text="Dodaj wiersz", command=self.add_blank_row).pack(side="left", padx=8)

        help_txt = ("Wybierz wiersz, edytuj pola i kliknij „Zastosuj”. "
                    "„Dodaj z kolumny…” tworzy wiersz mapujący istniejącą kolumnę. "
                    "„Dodaj wiersz” tworzy pusty wpis z transformacją 'constant'. "
                    "Transform 'format' używa 'sources' z presetu oraz szablonu w polu „Stała / Format”, "
                    "np. '{xcoord}, {ycoord}'. Brak wartości → ' - ' (poza 'constant').")
        ttk.Label(parent, text=help_txt, foreground="#555", wraplength=900, justify="left")\
            .pack(fill="x", padx=12, pady=(0,8))

        self.tree.bind("<<TreeviewSelect>>", self.on_select_row)

    def _build_opt_tab(self, parent):
        opt = ttk.LabelFrame(parent, text="Opcje dokumentu")
        opt.pack(fill="x", padx=8, pady=8)
        self.var_photo = tk.BooleanVar(value=True); self.var_photo_h = tk.DoubleVar(value=6.0)
        self.var_map = tk.BooleanVar(value=True);   self.var_map_h = tk.DoubleVar(value=6.0)
        self.var_break = tk.BooleanVar(value=True); self.var_comma = tk.BooleanVar(value=True)

        # Marginesy [cm] — DOMYŚLNE
        self.var_marg_l = tk.DoubleVar(value=2.0)
        self.var_marg_r = tk.DoubleVar(value=6.5)
        self.var_marg_t = tk.DoubleVar(value=3.9)
        self.var_marg_b = tk.DoubleVar(value=3.0)

        ttk.Checkbutton(opt, text="Dodaj ramkę na zdjęcie po każdym rekordzie", variable=self.var_photo)\
            .grid(row=0, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(opt, text="Wysokość [cm]:").grid(row=0, column=1, sticky="e")
        ttk.Entry(opt, textvariable=self.var_photo_h, width=6).grid(row=0, column=2, sticky="w", padx=6)

        ttk.Checkbutton(opt, text="Dziel ramkę zdjęcia na 2 pola (2 zdjęcia obok siebie)", variable=self.var_photo_split)\
            .grid(row=0, column=3, sticky="w", padx=12, pady=4)

        ttk.Checkbutton(opt, text="Dodaj ramkę na mapę (po zdjęciu/tabeli)", variable=self.var_map)\
            .grid(row=1, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(opt, text="Wysokość [cm]:").grid(row=1, column=1, sticky="e")
        ttk.Entry(opt, textvariable=self.var_map_h, width=6).grid(row=1, column=2, sticky="w", padx=6)

        ttk.Checkbutton(opt, text="Podział strony po każdym rekordzie", variable=self.var_break)\
            .grid(row=2, column=0, sticky="w", padx=8, pady=4)
        ttk.Checkbutton(opt, text="Użyj przecinka dziesiętnego", variable=self.var_comma)\
            .grid(row=2, column=1, sticky="w", padx=8, pady=4)

        # Marginesy [cm]
        ttk.Label(opt, text="Marginesy [cm]:").grid(row=3, column=0, sticky="w", padx=8, pady=(10,4))
        ttk.Label(opt, text="Lewy").grid(row=3, column=1, sticky="e");   ttk.Entry(opt, textvariable=self.var_marg_l, width=6).grid(row=3, column=2, sticky="w", padx=6)
        ttk.Label(opt, text="Prawy").grid(row=3, column=3, sticky="e");  ttk.Entry(opt, textvariable=self.var_marg_r, width=6).grid(row=3, column=4, sticky="w", padx=6)
        ttk.Label(opt, text="Górny").grid(row=4, column=1, sticky="e");  ttk.Entry(opt, textvariable=self.var_marg_t, width=6).grid(row=4, column=2, sticky="w", padx=6)
        ttk.Label(opt, text="Dolny").grid(row=4, column=3, sticky="e");  ttk.Entry(opt, textvariable=self.var_marg_b, width=6).grid(row=4, column=4, sticky="w", padx=6)

        # --- Sortowanie ---
        sortf = ttk.LabelFrame(parent, text="Sortowanie")
        sortf.pack(fill="x", padx=8, pady=(10,8))
        self.var_sort_col = tk.StringVar(value="")
        self.var_sort_asc = tk.BooleanVar(value=True)
        ttk.Label(sortf, text="Kolumna:").grid(row=0, column=0, sticky="e", padx=8, pady=6)
        self.cmb_sort_col = ttk.Combobox(sortf, textvariable=self.var_sort_col, width=40, values=[], state="readonly")
        self.cmb_sort_col.grid(row=0, column=1, sticky="w", padx=4, pady=6)
        ttk.Radiobutton(sortf, text="Rosnąco (A→Z / 0→9)", variable=self.var_sort_asc, value=True)\
            .grid(row=0, column=2, sticky="w", padx=12)
        ttk.Radiobutton(sortf, text="Malejąco (Z→A / 9→0)", variable=self.var_sort_asc, value=False)\
            .grid(row=0, column=3, sticky="w", padx=12)

        # --- Układ dokumentu ---
        lay = ttk.LabelFrame(parent, text="Układ dokumentu")
        lay.pack(fill="x", padx=8, pady=(10,8))
        self.var_layout = tk.StringVar(value="A")
        ttk.Radiobutton(lay, text="A — tabela, pod spodem: zdjęcie i/lub mapa", variable=self.var_layout, value="A")\
            .grid(row=0, column=0, sticky="w", padx=8, pady=6)
        ttk.Radiobutton(lay, text="B — tabela po lewej + scalona kolumna na zdjęcie po prawej; mapa pod spodem", variable=self.var_layout, value="B")\
            .grid(row=1, column=0, sticky="w", padx=8, pady=6)

        # --- CZCIONKA ---
        fontf = ttk.LabelFrame(parent, text="Czcionka")
        fontf.pack(fill="x", padx=8, pady=(10,8))
        ttk.Label(fontf, text="Rodzina:").grid(row=0, column=0, sticky="e", padx=8, pady=6)
        self.cmb_font = ttk.Combobox(fontf, textvariable=self.var_font_name, width=28,
                                     values=["Arial","Calibri","Times New Roman","Cambria","Verdana","Segoe UI","Tahoma","Garamond","Courier New"],
                                     state="readonly")
        self.cmb_font.grid(row=0, column=1, sticky="w", padx=4, pady=6)
        ttk.Label(fontf, text="Rozmiar [pt]:").grid(row=0, column=2, sticky="e", padx=8, pady=6)
        ttk.Entry(fontf, textvariable=self.var_font_size, width=6).grid(row=0, column=3, sticky="w", padx=4, pady=6)

        # --- SZEROKOŚCI KOLUMN (cm) ---
        colf = ttk.LabelFrame(parent, text="Szerokości kolumn [cm] (0 = AUTO)")
        colf.pack(fill="x", padx=8, pady=(10,8))
        ttk.Label(colf, text="Układ A — etykieta:").grid(row=0, column=0, sticky="e", padx=8, pady=6)
        ttk.Entry(colf, textvariable=self.var_A_label_cm, width=6).grid(row=0, column=1, sticky="w", padx=4, pady=6)
        ttk.Label(colf, text="Układ B — etykieta:").grid(row=1, column=0, sticky="e", padx=8, pady=6)
        ttk.Entry(colf, textvariable=self.var_B_label_cm, width=6).grid(row=1, column=1, sticky="w", padx=4, pady=6)
        ttk.Label(colf, text="wartość:").grid(row=1, column=2, sticky="e", padx=8, pady=6)
        ttk.Entry(colf, textvariable=self.var_B_value_cm, width=6).grid(row=1, column=3, sticky="w", padx=4, pady=6)
        ttk.Label(colf, text="zdjęcie:").grid(row=1, column=4, sticky="e", padx=8, pady=6)
        ttk.Entry(colf, textvariable=self.var_B_photo_cm, width=6).grid(row=1, column=5, sticky="w", padx=4, pady=6)

        # --- ZDJĘCIA: kompresja i skala ---
        photof = ttk.LabelFrame(parent, text="Zdjęcia — kompresja i skala")
        photof.pack(fill="x", padx=8, pady=(10,8))
        ttk.Label(photof, text="Jakość JPG (1–95):").grid(row=0, column=0, sticky="e", padx=8, pady=6)
        ttk.Entry(photof, textvariable=self.var_jpg_quality, width=6).grid(row=0, column=1, sticky="w", padx=4, pady=6)
        ttk.Label(photof, text="DPI eksportu (skalowanie):").grid(row=0, column=2, sticky="e", padx=8, pady=6)
        ttk.Entry(photof, textvariable=self.var_export_dpi, width=6).grid(row=0, column=3, sticky="w", padx=4, pady=6)
        ttk.Checkbutton(photof, text="Nie powiększaj małych zdjęć", variable=self.var_no_upscale).grid(row=0, column=4, sticky="w", padx=12, pady=6)

    # ---- Pomocnicze ----
    def _placeholder_in_extra(self, text):
        for w in self.extra_box.winfo_children(): w.destroy()
        ttk.Label(self.extra_box, text=text, foreground="#777", wraplength=900, justify="left")\
            .pack(fill="x", padx=12, pady=12)

    def rebuild_extra_columns(self):
        for w in self.extra_box.winfo_children(): w.destroy()
        self.extra_vars = {}
        if self.df is None:
            self._placeholder_in_extra("Wybierz plik Excel, aby rozpocząć."); return
        if not self.mapping:
            self._placeholder_in_extra("Wczytaj preset, aby zobaczyć kolumny pozostające poza mapą."); return
        mapped_sources = set()
        for m in self.mapping:
            if m.get("source"):
                mapped_sources.add(m["source"])
            for s in m.get("sources", []):
                mapped_sources.add(s)
        any_added = False
        for c in self.df.columns:
            if c in mapped_sources: continue
            v = tk.BooleanVar(value=False)
            ttk.Checkbutton(self.extra_box, text=c, variable=v).pack(anchor="w", padx=8, pady=2)
            self.extra_vars[c] = v; any_added = True
        if not any_added:
            self._placeholder_in_extra("Brak dodatkowych kolumn — preset obejmuje wszystkie źródła.")

    # ---- Presety (GUI) ----
    def load_preset_dialog(self):
        if self.df is None:
            messagebox.showwarning("Najpierw Excel", "Wczytaj plik Excel (żeby dopasować kolumny).")
            return
        p = filedialog.askopenfilename(initialdir="presets",
                                        filetypes=[("Preset JSON","*.json"),("All","*.*")])
        if not p: return
        try:
            data = json.loads(Path(p).read_text(encoding="utf-8"))
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie można wczytać presetu:\n{e}"); return
        aliases = data.get("aliases", {})
        new_mapping = []
        for item in data.get("mapping", []):
            m = dict(item)
            # resolve single source (if any)
            m["source"] = resolve_source_name(m.get("source",""), self.df.columns, aliases)
            # resolve multi-sources (if any)
            if isinstance(m.get("sources", None), list):
                m["sources"] = [resolve_source_name(s, self.df.columns, aliases) for s in m["sources"]]
            new_mapping.append(m)
        self.mapping = new_mapping
        opts = data.get("options", {})
        if opts:
            self.var_comma.set(bool(opts.get("decimal_comma", self.var_comma.get())))
            self.var_break.set(bool(opts.get("page_break", self.var_break.get())))
            ph = opts.get("photo", {})
            if "enabled" in ph: self.var_photo.set(bool(ph["enabled"]))
            if "height_cm" in ph: self.var_photo_h.set(float(ph["height_cm"]))
            mp = opts.get("map", {})
            if "enabled" in mp: self.var_map.set(bool(mp["enabled"]))
            if "height_cm" in mp: self.var_map_h.set(float(mp["height_cm"]))
        self.preset_label_var.set(data.get("name","(preset)"))
        self.refresh_tree()
        self.rebuild_extra_columns()
        self.cmb_source.configure(values=[""] + list(self.df.columns))

    # ---- Handlery standardowe ----
    def pick_excel(self):
        p = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx *.xls"), ("All","*.*")])
        if not p: return
        self.in_entry.delete(0, tk.END); self.in_entry.insert(0, p)
        self.load_columns()

    def pick_outdir(self):
        d = filedialog.askdirectory()
        if d: self.out_entry.delete(0, tk.END); self.out_entry.insert(0, d)

    def pick_imgdir(self):
        d = filedialog.askdirectory()
        if d: self.var_img_dir.set(d)

    def load_columns(self):
        path = self.in_entry.get().strip()
        if not path:
            messagebox.showwarning("Brak pliku","Wskaż plik Excel."); return
        try:
            self.df = read_excel_any(path)
        except Exception as e:
            messagebox.showerror("Błąd wczytywania", f"Nie udało się wczytać Excela:\n{e}\n\nJeśli to .xls, wymagany xlrd>=2.0.1.")
            return
        self.mapping = []
        self.preset_label_var.set("(brak)")
        self.refresh_tree()
        self.rebuild_extra_columns()
        cols_list = [""] + list(self.df.columns)
        self.cmb_source.configure(values=cols_list)
        self.cmb_sort_col.configure(values=cols_list)
        self.cmb_img_col.configure(values=cols_list)
        if "objektnummer" in self.df.columns:
            self.var_img_col.set("objektnummer")
        else:
            self.var_img_col.set(cols_list[0] if cols_list else "")
        self.var_sort_col.set("")

    def refresh_tree(self):
        self.tree.delete(*self.tree.get_children())
        for i, m in enumerate(self.mapping):
            if m.get("sources"):
                src_disp = ", ".join([s for s in m.get("sources", []) if s])
            else:
                src_disp = m.get("source","")
            self.tree.insert("", "end", iid=str(i),
                values=(i, "✓" if m.get("enabled",False) else "–",
                        m.get("label",""), src_disp,
                        m.get("transform","identity"), m.get("const","")))

    def on_select_row(self, *_):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0]); m = self.mapping[i]
        self.var_on.set(bool(m.get("enabled",False)))
        self.var_label.set(m.get("label",""))
        # przy multi-sources nie zmieniamy comboboxa — wyświetlamy pierwszy jako podpowiedź
        src0 = (m.get("sources", []) or [m.get("source","")])[0]
        self.var_source.set(src0)
        self.var_transform.set(m.get("transform","identity")); self.var_const.set(m.get("const",""))

    def apply_edit(self):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0])
        # zachowaj ewentualne źródła multi, jeśli transform pozostaje 'format'
        old_sources = self.mapping[i].get("sources", [])
        new_tf = self.var_transform.get().strip()
        new_map = {
            "enabled": bool(self.var_on.get()),
            "label": self.var_label.get().strip(),
            "source": self.var_source.get().strip(),
            "transform": new_tf,
            "const": self.var_const.get(),
        }
        if new_tf == "format":
            new_map["sources"] = old_sources  # edytujemy szablon/etykietę, źródła z presetu zostają
        self.mapping[i] = new_map
        self.refresh_tree(); self.tree.selection_set(str(i))
        self.rebuild_extra_columns()

    def move_selected(self, delta):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0]); j = i + delta
        if j < 0 or j >= len(self.mapping): return
        self.mapping[i], self.mapping[j] = self.mapping[j], self.mapping[i]
        self.refresh_tree(); self.tree.selection_set(str(j))

    def delete_selected(self):
        sel = self.tree.selection()
        if not sel: return
        i = int(sel[0]); del self.mapping[i]; self.refresh_tree()
        self.rebuild_extra_columns()

    def add_from_column(self):
        src = self.var_source.get().strip()
        if not src:
            messagebox.showinfo("Info","Wybierz źródło w polu 'Źródło', potem kliknij 'Dodaj z kolumny…'."); return
        self.mapping.append({"enabled": False, "label": src, "source": src, "transform":"identity", "const":"", "sources":[]})
        self.refresh_tree(); self.tree.selection_set(str(len(self.mapping)-1))
        self.rebuild_extra_columns()

    def add_blank_row(self):
        self.mapping.append({"enabled": True, "label": "", "source": "", "transform":"constant", "const":"", "sources":[]})
        self.refresh_tree(); self.tree.selection_set(str(len(self.mapping)-1))

    def run(self):
        if self.df is None:
            messagebox.showwarning("Brak danych","Najpierw wybierz plik Excel."); return
        out_dir = self.out_entry.get().strip() or os.path.dirname(self.in_entry.get()) or os.getcwd()
        if not os.path.isdir(out_dir):
            messagebox.showwarning("Folder","Podany folder wyjściowy nie istnieje."); return
        out_name = self.name_entry.get().strip() or "wynik.docx"
        extra = [k for k,v in self.extra_vars.items() if v.get()]
        if not self.mapping and not extra:
            messagebox.showwarning("Brak pól", "Mapa jest pusta i nie wybrano żadnych dodatkowych kolumn.\n"
                                               "Wczytaj preset lub dodaj wiersze/kolumny.")
            return

        logger = setup_logger(out_dir)
        try:
            logger.debug(f"[RUN] excel={self.in_entry.get().strip()}")
            logger.debug(f"[RUN] out_dir={out_dir}, out_name={out_name}")
            logger.debug(f"[RUN] images_dir={self.var_img_dir.get().strip()}, image_id_col={self.var_img_col.get().strip()}")
            logger.debug(f"[RUN] layout={self.var_layout.get()}, photo={self.var_photo.get()}, photo_h={self.var_photo_h.get()} cm, split={self.var_photo_split.get()}")
            logger.debug(f"[RUN] map={self.var_map.get()}, map_h={self.var_map_h.get()} cm")
            logger.debug(f"[RUN] margins=({self.var_marg_l.get()},{self.var_marg_r.get()},{self.var_marg_t.get()},{self.var_marg_b.get()}) cm")
            logger.debug(f"[RUN] font=({self.var_font_name.get()},{self.var_font_size.get()} pt)")
            logger.debug(f"[RUN] extra_cols={extra}")
            logger.debug(f"[RUN] mapping_rows={len(self.mapping)}")
            logger.debug(f"[RUN] jpg_quality={self.var_jpg_quality.get()}, export_dpi={self.var_export_dpi.get()}, no_upscale={self.var_no_upscale.get()}")

            # Sortowanie (opcjonalne)
            df_use = self.df
            col = getattr(self, "var_sort_col", tk.StringVar(value="")).get().strip() if hasattr(self, "var_sort_col") else ""
            if col:
                asc = getattr(self, "var_sort_asc", tk.BooleanVar(value=True)).get()
                def _sort_key(s):
                    if s.dtype == "O":
                        s_num = pd.to_numeric(s.astype(str).str.replace(",", ".", regex=False), errors="coerce")
                        if s_num.notna().any():
                            return s_num
                    return s
                try:
                    df_use = self.df.sort_values(by=col, ascending=asc, na_position="last", key=_sort_key)
                    logger.debug(f"[RUN] sorted by {col}, asc={asc}")
                except Exception as e:
                    logger.warning(f"[RUN] sort failed for {col}: {e}")
                    df_use = self.df

            out_file = build_docx(
                df_use,
                list(self.mapping), extra, out_dir, out_name,
                add_photo=self.var_photo.get(),
                photo_h_cm=float(self.var_photo_h.get()),
                photo_split=bool(self.var_photo_split.get()),
                add_map=self.var_map.get(),
                map_h_cm=float(self.var_map_h.get()),
                page_break=self.var_break.get(),
                use_decimal_comma=self.var_comma.get(),
                margin_left_cm=float(self.var_marg_l.get()),
                margin_right_cm=float(self.var_marg_r.get()),
                margin_top_cm=float(self.var_marg_t.get()),
                margin_bottom_cm=float(self.var_marg_b.get()),
                layout_mode=self.var_layout.get(),
                images_dir=self.var_img_dir.get().strip() or None,
                image_id_column=self.var_img_col.get().strip() or None,
                base_font_name=self.var_font_name.get().strip() or "Arial",
                base_font_size_pt=float(self.var_font_size.get() or 9.0),
                a_label_cm=float(self.var_A_label_cm.get() or 6.0),
                b_label_cm=float(self.var_B_label_cm.get() or 0.0),
                b_value_cm=float(self.var_B_value_cm.get() or 0.0),
                b_photo_cm=float(self.var_B_photo_cm.get() or 0.0),
                jpg_quality=int(self.var_jpg_quality.get() or 90),
                export_dpi=int(self.var_export_dpi.get() or 450),
                no_upscale=bool(self.var_no_upscale.get()),
                logger=logger,
            )
        except Exception as e:
            logger.error(f"[RUN] generation failed: {e}")
            logger.debug(traceback.format_exc())
            messagebox.showerror("Błąd", f"Nie udało się wygenerować DOCX:\n{e}\n\nSzczegóły w logu:\n{logger.log_path}")
            return
        messagebox.showinfo("Gotowe", f"Zapisano:\n{out_file}\n\nLog:\n{logger.log_path}")

if __name__ == "__main__":
    App().mainloop()